import base64
import io
import logging
import mimetypes
from fastapi import UploadFile, HTTPException

log = logging.getLogger(__name__)

ALLOWED_MIME_TYPES = {
    "application/pdf",
    "image/jpeg",
    "image/png",
    "image/webp",
    "image/heic",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 MB
MAX_TEXT_CHARS = 60_000

DOCX_MIME_TYPES = {
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _extract_docx_text(content: bytes) -> str:
    log.debug("Extracting text from DOCX (%d bytes)", len(content))
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n\n".join(paragraphs)
        log.debug("DOCX extraction complete: %d paragraphs, %d chars", len(paragraphs), len(text))
        return text
    except Exception as e:
        log.error("DOCX extraction failed: %s", e, exc_info=True)
        raise HTTPException(status_code=422, detail=f"Could not read Word document: {e}")


def extract_pdf_text(content: bytes) -> str:
    log.debug("Extracting text layer from PDF (%d bytes)", len(content))
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Page {i + 1}]\n{text.strip()}")
        result = "\n\n".join(pages)
        log.debug(
            "PDF text extraction: %d/%d pages had text, %d total chars",
            len(pages), len(reader.pages), len(result),
        )
        return result
    except Exception as e:
        log.warning("PDF text extraction failed (visual-only doc?): %s", e)
        return ""


async def process_upload(file: UploadFile) -> tuple[str, str, str]:
    """
    Returns (payload, mime_type, raw_text).
    - Binary files (PDF/image): payload=base64, raw_text=extracted text (empty for images).
    - Word docs: payload=extracted text, mime_type='text/plain', raw_text=same.
    """
    log.info("Processing upload: filename=%r content_type=%r", file.filename, file.content_type)

    content = await file.read()
    size_kb = len(content) / 1024
    log.debug("Read %d bytes (%.1f KB) from %r", len(content), size_kb, file.filename)

    if len(content) > MAX_FILE_SIZE:
        log.warning("File too large: %d bytes from %r", len(content), file.filename)
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 100MB.")

    mime_type = file.content_type or mimetypes.guess_type(file.filename or "")[0] or ""
    log.debug("Detected mime_type=%r for %r", mime_type, file.filename)

    filename = (file.filename or "").lower()
    if mime_type == "application/octet-stream":
        if filename.endswith(".docx"):
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif filename.endswith(".doc"):
            mime_type = "application/msword"
        log.debug("Normalised octet-stream to mime_type=%r based on extension", mime_type)

    if mime_type not in ALLOWED_MIME_TYPES:
        log.warning("Rejected unsupported mime_type=%r for %r", mime_type, file.filename)
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {mime_type}. Accepted: PDF, JPEG, PNG, WEBP, HEIC, DOC, DOCX."
        )

    if mime_type in DOCX_MIME_TYPES:
        text = _extract_docx_text(content)
        truncated = text[:MAX_TEXT_CHARS]
        if len(text) > MAX_TEXT_CHARS:
            log.info("DOCX text truncated from %d to %d chars", len(text), MAX_TEXT_CHARS)
        log.info("DOCX processed: %r -> %d chars of text", file.filename, len(truncated))
        return truncated, "text/plain", truncated

    raw_text = ""
    if mime_type == "application/pdf":
        raw_text = extract_pdf_text(content)[:MAX_TEXT_CHARS]
        log.info("PDF processed: %r -> %d chars of text extracted", file.filename, len(raw_text))
    else:
        log.info("Image processed: %r mime=%r (no text extraction)", file.filename, mime_type)

    b64 = base64.b64encode(content).decode("utf-8")
    log.debug("Base64 payload: %d chars for %r", len(b64), file.filename)
    return b64, mime_type, raw_text
